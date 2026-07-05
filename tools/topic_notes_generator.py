from __future__ import annotations

import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/arnavmalhotra/IdeaProjects/helloworld/src")
OUT = Path("/private/tmp/helloworld-topic-notes")

SKIP_DIRS = {
    "heap-notes",
    "main/resources",
}


@dataclass
class Note:
    topic: str
    class_name: str
    overview: str
    intuition: str
    brute_force: str
    bottleneck: str
    better_solution: str
    complexity: str
    teaching_point: str
    indicators: str
    diagram: str
    refresher: str


def normalize(name: str) -> str:
    return re.sub(r"(?<!^)(?=[A-Z])", " ", name).replace("_", " ").strip()


def topic_overview(topic: str) -> str:
    mapping = {
        "heap": "Heap mechanics, priority ordering, top-k queries, and heap transformation patterns.",
        "priorityQueue": "Priority-queue style selection problems, streaming top-k, and min/max heap usage.",
        "binaryTree": "Tree recursion, traversal patterns, path problems, BST reasoning, and tree transformation.",
        "graph": "Traversal, shortest path, minimum spanning tree, cycle detection, and grid-to-graph patterns.",
        "stack": "Monotonic stack, parsing, expression-like problems, and stack-based state compression.",
        "linkedList": "Pointer manipulation, reversals, cycle handling, merges, and list rearrangement.",
        "arrays": "Array transformation, prefix-sum, two-pointer, greedy, backtracking, and DP-style reasoning.",
        "leetCode": "Mixed interview practice spanning arrays, strings, DP, graph, stack, and greedy patterns.",
        "dynamic": "Dynamic programming fundamentals and optimization by storing subproblem results.",
        "sort": "Comparison and non-comparison sorting, partitioning, and order-statistic thinking.",
        "trie": "Prefix-tree based string matching and word segmentation problems.",
        "Thread": "Concurrency, coordination, synchronization, and lock-based sequencing.",
        "builder": "Builder pattern and object construction with readability and immutability in mind.",
        "singleton": "Singleton pattern variants and lifecycle control.",
        "Serialization": "Object persistence and custom serialization concerns.",
        "streams": "Functional-style transformations and stream pipelines.",
        "consumer_supplier": "Functional interface usage and callback-style APIs.",
        "AnonymousClass": "Anonymous classes, interfaces, and small behavior overrides.",
        "Queue": "Queue fundamentals and linear FIFO data flow.",
        "Number_Theory": "Digit math and elementary number-theory style manipulation.",
        "parking_lot": "Low-level system design, entities, and object composition.",
        "main": "Mixed Java practice: inheritance, interfaces, recursion, and miscellaneous interview demos.",
    }
    return mapping.get(topic, "General Java practice and concept exploration.")


def topic_diagram(topic: str) -> str:
    diagram_map = {
        "heap": "problem -> build heap -> fix root/children -> repeat -> answer",
        "binaryTree": "tree shape -> recursive relation -> traversal/state -> answer",
        "graph": "nodes/edges -> traversal or relaxation -> visited/dist -> answer",
        "stack": "incoming sequence -> push/pop rule -> monotonic state -> answer",
        "arrays": "raw array -> pattern detect -> structure/greedy/DP -> answer",
        "linkedList": "nodes -> pointer rewiring -> local fix -> whole list result",
        "trie": "word/prefix -> trie path -> pruning/memo -> answer",
        "dynamic": "state definition -> transition -> base case -> table/answer",
    }
    return diagram_map.get(topic, "problem -> identify pattern -> choose structure -> refine -> answer")


def contains_any(text: str, words: Iterable[str]) -> bool:
    return any(word in text for word in words)


def generic_note(topic: str, class_name: str) -> Note:
    readable = normalize(class_name)
    overview = f"{readable} is a {topic} practice problem that teaches how to recognize the core pattern instead of forcing a brute-force search."
    intuition = "Start by asking what repeated work the naive solution would do, then look for the invariant or data structure that removes that repetition."
    brute = "The naive route usually recomputes from scratch, scans too much, or explores every option without remembering previous work."
    bottleneck = "That becomes slow because the same subproblem, traversal, or comparison is repeated many times."
    better = "The improved approach keeps just enough state to avoid repetition, using the appropriate structure for this topic."
    complexity = "Brute force is typically quadratic or worse; the optimized version is usually linear, linearithmic, or O(log n) per update depending on the pattern."
    teaching = f"The goal of {readable} is to train the eye to spot the topic-level pattern quickly."
    indicators = f"Look for keywords like {topic.lower()}, repeated state, ordering, prefix checks, traversal, or local-to-global reasoning."
    refresher = f"{readable}: identify the repeated work, choose the topic structure, and reduce the state until the answer becomes direct."
    return Note(topic, class_name, overview, intuition, brute, bottleneck, better, complexity, teaching, indicators, topic_diagram(topic), refresher)


def classify_heap(name: str) -> Note:
    n = name.lower()
    if "drivecar" in n:
        return Note("heap", name,
                    "This class is about extracting the needed maximum excess over a threshold instead of overusing heap machinery.",
                    "Ask whether a single scan can track the best candidate directly.",
                    "A naive approach would sort or build a heap even though the answer is only one aggregate value.",
                    "Sorting or heap construction adds complexity without improving the final answer for a one-shot query.",
                    "Use one pass and keep the maximum value above the threshold.",
                    "Brute force: O(n log n) if you sort; optimized scan: O(n) time and O(1) space.",
                    "The lesson is to prefer direct aggregation when the output is a single statistic.",
                    "Single-output query, threshold comparison, and no need to preserve full ordering.",
                    "problem -> scan once -> track max excess -> return answer",
                    "DriveCar: one scan is enough; avoid heap or sort when the goal is just a max excess.")
    if "heapify" in n:
        return Note("heap", name,
                    "Heapify teaches how to restore heap order from a subtree and then use that to build a heap bottom-up.",
                    "Start with parent-child order and fix local violations from the last non-leaf node upward.",
                    "A naive route would repeatedly sort or bubble values into place one by one.",
                    "That repeats too much work because each insertion would reprocess the same levels.",
                    "Bottom-up heapify fixes the root of each subtree once its children are already valid heaps.",
                    "Build heap: O(n); heap sort: O(n log n); space: O(1) in-place.",
                    "The lesson is that local fixes composed bottom-up can produce a globally correct heap efficiently.",
                    "Array-based tree layout, repeated root extraction, and kth/top-k style queries.",
                    "tree array -> heapify -> build heap -> repeated root swap -> answer",
                    "Heapify: fix children first, then parent; that is why build-heap is linear.")
    if "insertanddelete" in n:
        return Note("heap", name,
                    "This class is about maintaining a max-heap under insertion and deletion.",
                    "Insert by appending then bubbling up; delete by moving the last value to the root and bubbling down.",
                    "A naive route would rebuild the whole heap after every update.",
                    "Rebuilding wastes time because only one path in the tree changes.",
                    "Bubble-up and bubble-down touch only the affected path from leaf to root or root to leaf.",
                    "Each operation is O(log n); array resizing can add extra copying if the container grows.",
                    "The lesson is that heaps are dynamic priority structures, not sorted arrays.",
                    "Frequent top-priority updates, insertions, or removals.",
                    "insert -> bubble up; delete root -> bubble down; repeat",
                    "Insert/Delete: only the affected path changes, so update in logarithmic time.")
    if "mintomaxheap" in n:
        return Note("heap", name,
                    "This class teaches heap conversion: turning a min-heap-shaped array into a max-heap.",
                    "The key is to fix every internal node bottom-up, not just swap the largest value to the top.",
                    "A naive route would move the maximum to the root once and stop.",
                    "That ignores violations deeper in the tree, so the structure is still not a valid max-heap.",
                    "Run max-heapify from the last internal node to the root.",
                    "Bottom-up conversion is O(n) and in-place.",
                    "The lesson is that heap validity is global; one visible fix is not enough.",
                    "Need full heap property, subtree correctness, and parent-child ordering.",
                    "scan internal nodes -> heapify each subtree -> max-heap",
                    "MinToMaxHeap: global validity comes from fixing all subtrees, not just the root.")
    return generic_note("heap", name)


def classify_binary_tree(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["bfs_dfs", "levelorder", "zigzag", "traversal"]):
        return Note("binaryTree", name,
                    "This class is about tree traversal styles: DFS, BFS, level order, and zigzag variants.",
                    "Start with recursive DFS or queue-based BFS depending on the traversal order needed.",
                    "A naive approach often re-traverses the tree for each level or each view.",
                    "That repeats work and makes the code harder to generalize.",
                    "Use one traversal pass and store just the state needed for the chosen order.",
                    "Traversal is O(n); extra space is O(h) for recursion or O(w) for breadth-first levels.",
                    "The lesson is to match traversal strategy to the output shape.",
                    "level order / DFS / view problems / alternating direction output.",
                    "tree -> choose DFS or BFS -> maintain small state -> answer",
                    f"{normalize(name)}: the main choice is DFS versus BFS, then carrying just enough state.")
    if contains_any(n, ["pathsum3", "pathsum2", "pathsum"]):
        return Note("binaryTree", name,
                    "This class teaches how path-sum problems differ: exact path, any downward path, or count of matching paths.",
                    "Think in terms of accumulated sums along the current path and whether you need counting or existence.",
                    "A naive route would enumerate every path and re-sum it repeatedly.",
                    "That can become quadratic or worse across all root-to-node combinations.",
                    "Use recursion with running sums, and use prefix sums when counting many matching paths efficiently.",
                    "Brute force: O(n^2) to O(n^3) depending on path counting; optimized prefix-sum DFS: O(n).",
                    "The lesson is to decide whether you need a boolean, a path, or a count before choosing the recursion state.",
                    "path accumulation, backtracking, and subtree prefix counts.",
                    "recursive path -> accumulate sum -> backtrack -> count/return",
                    f"{normalize(name)}: prefix state turns repeated path checks into one DFS pass.")
    if contains_any(n, ["maxpathsum", "maxsumofroottoleafpaths", "tilt", "diameter", "deepestleavessum"]):
        return Note("binaryTree", name,
                    "This class is about aggregating a tree value from children and combining local results into a global answer.",
                    "Think bottom-up: each node returns a contribution, while the global best is updated at each visit.",
                    "A naive route would recompute subtree information for every node.",
                    "That duplicates subtree work and often leads to O(n^2) behaviour.",
                    "Use DFS that returns the best downward contribution and updates a global tracker for the overall optimum.",
                    "Most of these are O(n) time and O(h) space.",
                    "The lesson is that tree DP often means 'return one value upward, keep another value globally.'",
                    "subtree contribution, global optimum, and recursive aggregation.",
                    "node -> combine children -> return contribution -> update global best",
                    f"{normalize(name)}: return one number to the parent and keep the true answer separately.")
    if contains_any(n, ["lca", "lowestcommonancestor", "cousin", "same", "subtree", "issametree"]):
        return Note("binaryTree", name,
                    "This class teaches recursive comparison and ancestor reasoning in trees.",
                    "Ask whether the answer exists in the left subtree, the right subtree, or at the current node.",
                    "A naive route would search path lists for both nodes and compare them fully.",
                    "That is workable but often more verbose than necessary.",
                    "Use recursion to let each child report whether it contains a target, and combine the answers at the current node.",
                    "Most solutions are O(n) time and O(h) space.",
                    "The lesson is that tree ancestry can usually be solved by letting recursion report presence upward.",
                    "ancestor discovery, subtree checks, and path comparison.",
                    "search left/right -> merge return values -> decide at current node",
                    f"{normalize(name)}: recursion answers where the targets are, so the split point becomes obvious.")
    if contains_any(n, ["invert", "flatten", "burn", "sortedarraytobst", "convertsortedarraytobst", "constructtreefromviews", "view", "swappednode", "min_depth", "height", "size", "mindepth", "min_absolute_diff"]):
        return Note("binaryTree", name,
                    "This class is a tree transformation or structural utility problem.",
                    "Work out whether the operation is best handled top-down, bottom-up, or level-by-level.",
                    "A naive route often copies nodes into another structure or rebuilds the whole tree repeatedly.",
                    "That loses the advantage of the original shape and often adds unnecessary memory use.",
                    "Use recursive swap/relink, inorder reasoning for BST properties, or BFS for level-based tasks.",
                    "Most transformations are O(n) time; extra space depends on recursion or queue usage.",
                    "The lesson is to treat the tree as a structure to be rewired, not flattened into arrays unless required.",
                    "rewiring nodes, BST order, level traversal, and recursive mutation.",
                    "tree shape -> choose recursion or BFS -> rewire links -> answer",
                    f"{normalize(name)}: the key is whether the shape should change locally or level-wise.")
    return generic_note("binaryTree", name)


def classify_graph(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["dfs", "bfs", "pathexistence", "numberofislands", "snake"]):
        return Note("graph", name,
                    "This class is about graph/grid traversal with visited-state management.",
                    "Start from one node or cell and explore reachable neighbors while marking visited.",
                    "A naive route would try every route independently or revisit cells repeatedly.",
                    "That creates exponential explosion or unnecessary repeated work.",
                    "Use BFS/DFS with a visited set or grid marks to process each node once.",
                    "Traversal is O(V + E) for graphs or O(rows * cols) for grids.",
                    "The lesson is to detect reachability and connected components with one systematic traversal.",
                    "reachability, connected components, and grid flood fill.",
                    "start node -> visit neighbors -> mark visited -> finish",
                    f"{normalize(name)}: traversal with a visited set prevents repeated exploration.")
    if contains_any(n, ["cycle"]):
        return Note("graph", name,
                    "This class teaches cycle detection in a graph.",
                    "Use DFS state or parent tracking to decide whether you are re-entering an active path.",
                    "A naive route would keep exploring without remembering the current recursion path.",
                    "Without path state, you cannot distinguish a back-edge from a fresh edge.",
                    "Use visited plus recursion-stack logic for directed graphs, or parent tracking for undirected graphs.",
                    "Typical complexity is O(V + E).",
                    "The lesson is that cycle detection is about path state, not just global visited status.",
                    "active path, back-edge, and visited-state discipline.",
                    "DFS -> track active path -> detect revisit -> answer",
                    f"{normalize(name)}: cycles are found by distinguishing 'seen before' from 'seen on current path'.")
    if contains_any(n, ["dijkstra", "bellmanford", "floydwarshall"]):
        return Note("graph", name,
                    "This class is about shortest-path reasoning under different edge constraints.",
                    "Choose the algorithm based on whether edges are weighted, negative, or all-pairs.",
                    "A naive route would enumerate all paths and compare lengths.",
                    "That is too slow once the graph grows because path count explodes.",
                    "Use Dijkstra for non-negative weights, Bellman-Ford for negative edges, or Floyd-Warshall for all-pairs small graphs.",
                    "Dijkstra: O((V+E) log V); Bellman-Ford: O(VE); Floyd-Warshall: O(V^3).",
                    "The lesson is to match the shortest-path method to the weight rules.",
                    "weighted edges, relaxation, and choosing the right shortest-path family.",
                    "edges -> relax distances -> frontier/DP -> answer",
                    f"{normalize(name)}: the edge constraints decide the algorithm, not the node count alone.")
    if contains_any(n, ["prim", "kruskal", "mst"]):
        return Note("graph", name,
                    "This class teaches minimum spanning tree construction.",
                    "Use edge selection that always keeps the partial structure acyclic and cheap.",
                    "A naive route would try every spanning tree.",
                    "That is combinatorially impossible beyond small graphs.",
                    "Use Prim or Kruskal with greedy selection and cycle avoidance.",
                    "Prim and Kruskal are both near O(E log V) with the right data structures.",
                    "The lesson is that MST problems are greedy because each safe local choice preserves global optimality.",
                    "edge ordering, greedy cut property, and cycle prevention.",
                    "choose cheapest safe edge -> grow tree -> repeat",
                    f"{normalize(name)}: greedy edge selection is enough when the cut property holds.")
    return generic_note("graph", name)


def classify_stack(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["nextgreater", "stockspanner", "slidingwindowmaximum", "largesthistogram", "largestrectangle"]):
        return Note("stack", name,
                    "This class teaches the monotonic stack pattern.",
                    "Keep elements in sorted-ish stack order so each new value can remove stale candidates quickly.",
                    "A naive route would compare each item with many future or past elements.",
                    "That leads to O(n^2) scans for next-greater or window maximum style questions.",
                    "Use a monotonic stack or deque so each element is pushed and popped at most once.",
                    "The optimized pattern is O(n) time and O(n) space.",
                    "The lesson is that one-pass state compression beats repeated local comparisons.",
                    "nearest greater/smaller, windows, and area-from-heights reasoning.",
                    "scan -> maintain monotonic stack -> resolve pending elements -> answer",
                    f"{normalize(name)}: monotonic stacks turn repeated comparisons into one linear pass.")
    if contains_any(n, ["balancedparenthesis", "stringdecode", "simplifypath", "removekdigits", "asteroidcollision", "celebrity"]):
        return Note("stack", name,
                    "This class is about using a stack to model parsing, cancellation, or directional collision.",
                    "Push state while scanning; pop when the new token invalidates the top state.",
                    "A naive route would simulate the whole process with repeated rescans after each change.",
                    "That creates unnecessary reprocessing and messy code.",
                    "Use a stack to store the current valid prefix or unresolved items.",
                    "Most solutions are O(n) time and O(n) space.",
                    "The lesson is that stacks are good for 'last unresolved thing first' logic.",
                    "parentheses, decoding, path simplification, and collision resolution.",
                    "scan tokens -> push/pop based on rule -> final stack state",
                    f"{normalize(name)}: stack state is the natural fit for nested or cancelling operations.")
    return generic_note("stack", name)


def classify_linked_list(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["reverse", "rotate", "palindrom", "remove", "create", "cycle", "intersection", "merge"]):
        return Note("linkedList", name,
                    "This class is about pointer rewiring rather than value manipulation.",
                    "Think in terms of local pointer changes while preserving the rest of the chain.",
                    "A naive route would copy nodes into arrays or rebuild lists from scratch.",
                    "That wastes memory and hides the actual pointer logic.",
                    "Use fast/slow pointers, dummy nodes, sublist reversal, or merge-style splicing as needed.",
                    "Most linked-list operations are O(n) time and O(1) extra space.",
                    "The lesson is that linked lists reward careful pointer discipline.",
                    "pointer movement, local rewiring, and safe boundary handling.",
                    "list -> find pivot -> rewire pointers -> continue",
                    f"{normalize(name)}: the answer is in pointer updates, not in value sorting.")
    return generic_note("linkedList", name)


def classify_arrays_like(folder: str, name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["sort", "selectionsort", "mergesort", "quicksort", "countsort", "radixsort", "insertionsort"]):
        return Note(folder, name,
                    "This class teaches sorting or partitioning as a way to impose order.",
                    "Start from direct comparison sorting and then compare it with a more specialized linear or divide-and-conquer method.",
                    "A naive route is to repeatedly place each element into its correct position with too much scanning.",
                    "That may become quadratic or requires repeated passes over the data.",
                    "Use the exact sorting strategy the data allows: partitioning, counting, radix, or merge-based order.",
                    "Complexity depends on the chosen sort: O(n^2), O(n log n), or O(n + k) for constrained domains.",
                    "The lesson is to choose a sort based on data structure and constraints, not habit.",
                    "ordering, partitioning, stability, and domain-specific optimization.",
                    "array -> choose sort strategy -> partition/count/merge -> answer",
                    f"{normalize(name)}: the key question is which sort matches the constraints.")
    if contains_any(n, ["binarysearch", "rotated", "sortedinfinite", "findmin", "minimum", "maxavg", "hindex", "koko", "search"]):
        return Note(folder, name,
                    "This class is about searching with structure rather than checking every element one by one.",
                    "Look for monotonicity, sortedness, or a decision boundary that lets you discard half the search space.",
                    "A naive route would scan all values or all positions.",
                    "That ignores the fact that the data already has an order or a threshold property.",
                    "Use binary search or a binary-search-like feasibility check when the answer space is ordered.",
                    "Usually O(log n) for direct search or O(log range) with a check function.",
                    "The lesson is to ask whether the problem has a monotonic predicate.",
                    "sorted order, monotonic checks, and half-space elimination.",
                    "search space -> compare middle -> eliminate half -> answer",
                    f"{normalize(name)}: monotonicity is the signal that binary search applies.")
    if contains_any(n, ["subarray", "window", "sliding", "prefix", "equilibrium", "triplets", "sum", "subsequence", "subsets", "windowmaximum", "maxsum", "minsum", "leftelementsmall"]):
        return Note(folder, name,
                    "This class is about scanning contiguous regions or tracking cumulative state.",
                    "Use prefix sums, sliding windows, or two pointers depending on whether the window size changes.",
                    "A naive route would examine every subarray directly.",
                    "That quickly becomes O(n^2) or worse as the number of candidate ranges grows.",
                    "Use a running total, prefix map, or window adjustment to avoid recomputing sums from scratch.",
                    "Many such problems improve from O(n^2) to O(n).",
                    "The lesson is to recognize when a range query can be turned into incremental state.",
                    "contiguous ranges, cumulative totals, and controlled movement of endpoints.",
                    "range -> maintain running state -> move boundaries -> answer",
                    f"{normalize(name)}: contiguous range problems usually reward prefix or window state.")
    if contains_any(n, ["kadane", "maxsumsubarray", "complement", "rainwater", "water", "container", "maxwater", "product", "leader", "missing", "majority"]):
        return Note(folder, name,
                    "This class teaches a classic array trick: local state, greedy choice, or a small invariant that beats full enumeration.",
                    "Ask whether the answer can be updated as you scan once from left to right or from both ends.",
                    "A naive route would try every split, every pair, or every candidate region.",
                    "That wastes time because the answer can often be maintained incrementally.",
                    "Use Kadane, two pointers, prefix/suffix extrema, or a voting invariant depending on the exact pattern.",
                    "Brute force is often O(n^2); the optimized scan is O(n).",
                    "The lesson is to reduce the problem to an invariant that can survive one pass.",
                    "greedy scan, local best, and global accumulator.",
                    "scan once -> keep invariant -> update answer -> finish",
                    f"{normalize(name)}: the trick is usually an invariant that survives a single sweep.")
    if contains_any(n, ["nqueen", "permutation", "subset", "combination", "backtracking", "maze", "ratinamaze", "palindromepartitioning", "generateparentheses", "findallpermutations"]):
        return Note(folder, name,
                    "This class is a backtracking problem: build choices recursively and undo them when they fail.",
                    "Try one choice at a time, then revert the state and try the next option.",
                    "A naive route would brute-force all possibilities without pruning.",
                    "That explodes combinatorially and becomes hard to manage.",
                    "Use recursion with pruning and explicit undo steps to explore only feasible branches.",
                    "These are generally exponential in the worst case, but pruning makes them practical.",
                    "The lesson is to recognize decision trees and control the explosion with constraints.",
                    "choice tree, pruning, and recursive undo.",
                    "choose -> recurse -> undo -> next option",
                    f"{normalize(name)}: recursive search with backtracking is the natural fit.")
    if contains_any(n, ["stock", "job", "booksallocation", "platform", "interval", "overlapping", "canplaceflower", "gasstation"]):
        return Note(folder, name,
                    "This class is about greedy decisions and feasibility checking.",
                    "Look for a local choice that stays globally safe, or a feasibility test that can be binary-searched.",
                    "A naive route would try many arrangements or simulate all start points exhaustively.",
                    "That is too expensive when the greedy invariant already tells us what can work.",
                    "Use a greedy rule, a sorted order, or a feasibility check depending on the problem statement.",
                    "Usually O(n log n) due to sorting or O(n) for a single greedy scan.",
                    "The lesson is to ask whether the problem is about proving a safe local choice.",
                    "greedy feasibility, safe choice, and exchange argument.",
                    "sort/scan -> keep feasible state -> pick safe choice -> answer",
                    f"{normalize(name)}: greedy works when a safe local decision stays optimal.")
    if contains_any(n, ["matrix", "spirally", "zigzag", "reversearray", "merge", "productmatrix", "gameoflife", "slidingwindowmaximum"]):
        return Note(folder, name,
                    "This class teaches a matrix or sequence manipulation pattern with precise iteration order.",
                    "Decide whether the answer needs row-wise, column-wise, spiral, or window-based state.",
                    "A naive route would rebuild the whole answer from repeated scans or copies.",
                    "That adds avoidable extra passes and memory churn.",
                    "Use the exact traversal or transformation order the problem asks for, often in-place.",
                    "Usually O(n) or O(rows*cols) with minimal extra space.",
                    "The lesson is to respect the required traversal order instead of reshaping data unnecessarily.",
                    "order of traversal, in-place update, and shape transformation.",
                    "iterate in required order -> update in place -> finish",
                    f"{normalize(name)}: matrix/sequence problems often hinge on traversal order.")
    return generic_note(folder, name)


def classify_dynamic(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["knapsack", "mincoins", "uglynumbers"]):
        return Note("dynamic", name,
                    "This class is about classic dynamic programming state building.",
                    "Start with subproblem definitions and think about whether the answer depends on the previous row, previous coin, or previous sum.",
                    "A naive route would recurse into every combination repeatedly.",
                    "That repeats the same states again and again.",
                    "Store solved subproblems in a table or memoized recursion.",
                    "Typical complexity improves from exponential to polynomial, often O(n*amount) or similar.",
                    "The lesson is to define the state before you define the code.",
                    "subproblem state, memoization, and transition design.",
                    "state -> transition -> memo table -> answer",
                    f"{normalize(name)}: DP starts with the state, not with loops.")
    return generic_note("dynamic", name)


def classify_trie(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["insertandsearch", "trie"]):
        return Note("trie", name,
                    "This class teaches the core trie structure for prefix matching.",
                    "Store characters along a path so common prefixes are shared.",
                    "A naive route would scan every word for every query.",
                    "That repeats the same prefix comparisons many times.",
                    "Use a trie to share prefix work and make lookup proportional to word length.",
                    "Insert/search become O(length of word).",
                    "The lesson is to trade memory for repeated-prefix speed.",
                    "prefix sharing, exact match, and incremental character traversal.",
                    "word -> follow characters -> prefix node -> answer",
                    f"{normalize(name)}: tries compress repeated prefixes.")
    if contains_any(n, ["wordbreak"]):
        return Note("trie", name,
                    "This class is about segmentation with prefix reuse.",
                    "Ask whether each substring can be validated from the current prefix state.",
                    "A naive route would try every split point recursively without remembering failures.",
                    "That creates repeated work and exponential branching.",
                    "Use trie lookups plus memoization or DP over split positions.",
                    "Often O(n^2) or better with memoization and trie pruning.",
                    "The lesson is that prefix data structures pair naturally with word segmentation.",
                    "prefix matching, segmentation, and memoized recursion.",
                    "string -> split points -> prefix check -> answer",
                    f"{normalize(name)}: word-break is prefix DP with shared trie lookups.")
    return generic_note("trie", name)


def classify_priority_queue(name: str) -> Note:
    n = name.lower()
    if contains_any(n, ["kthlargest", "findmedian", "topkfrequent", "connectropes"]):
        return Note("priorityQueue", name,
                    "This class teaches how a heap-backed priority queue solves repeated top-element selection.",
                    "Keep only the needed top candidates instead of sorting everything.",
                    "A naive route would sort the entire input or scan it again and again.",
                    "That wastes work when only a small priority frontier matters.",
                    "Use a min-heap or max-heap depending on whether you need kth largest, median, or cheapest merge cost.",
                    "Typical complexity is O(n log k) or O(n log n) depending on heap size.",
                    "The lesson is to store only the frontier of relevance.",
                    "top-k, running median, and merge-cost minimization.",
                    "stream -> heap frontier -> extract best -> continue",
                    f"{normalize(name)}: priority queues are for 'best candidate now' problems.")
    return generic_note("priorityQueue", name)


def classify_thread(name: str) -> Note:
    n = name.lower()
    return Note("Thread", name,
                "This class is about concurrency control, sequencing, and synchronization.",
                "Decide whether the problem needs mutual exclusion, ordering, or waiting for a condition.",
                "A naive route would let threads run independently and hope the output order is correct.",
                "That can cause race conditions or nondeterministic behaviour.",
                "Use locks, reentrant locks, or coordination primitives to define the allowed interleaving.",
                "Complexity is usually dominated by synchronization behaviour rather than pure asymptotic cost.",
                "The lesson is that concurrency problems are about state ownership and ordering guarantees.",
                "shared state, ordering, and race prevention.",
                "threads -> acquire lock/turn -> act -> release -> next",
                f"{normalize(name)}: thread problems teach safe coordination, not raw speed.")


def classify_basic(topic: str, name: str) -> Note:
    if topic == "builder":
        return Note(topic, name,
                    "This class teaches object construction with the builder pattern.",
                    "Ask whether constructing the object in one call is hard to read or easy to misuse.",
                    "A naive route would expose a telescoping constructor with many parameters.",
                    "That becomes hard to read and brittle when optional fields grow.",
                    "Use a builder to stage construction and make intent explicit.",
                    "Construction remains O(1); the win is readability and correctness.",
                    "The lesson is to separate object creation from object use.",
                    "optional fields, fluent API, and readable construction.",
                    "config -> builder -> build() -> object",
                    f"{normalize(name)}: builder reduces constructor noise and improves clarity.")
    if topic == "singleton":
        return Note(topic, name,
                    "This class teaches how to control instance creation so only one object exists.",
                    "Ask when and how the instance should be created and shared.",
                    "A naive route would expose public construction everywhere.",
                    "That allows multiple objects and breaks the singleton intent.",
                    "Use eager or static initialization to guarantee one shared instance.",
                    "Instance access is O(1).",
                    "The lesson is about lifecycle control and shared identity.",
                    "single instance, shared access, and lifecycle constraints.",
                    "request -> access singleton -> reuse instance",
                    f"{normalize(name)}: singleton means one controlled instance, not many copies.")
    if topic == "Serialization":
        return Note(topic, name,
                    "This class teaches converting objects to a storable or transferable form and back again.",
                    "Think about which fields must persist and which should be reconstructed.",
                    "A naive route would rely on default behavior without checking compatibility or custom rules.",
                    "That can break on version changes or skip important invariants.",
                    "Use explicit serialization hooks or default serialization carefully.",
                    "Cost depends on object size and I/O, typically linear in serialized data length.",
                    "The lesson is to treat object persistence as a contract, not a side effect.",
                    "object state, persistence, and reconstruction.",
                    "object -> serialize -> store/transmit -> deserialize",
                    f"{normalize(name)}: serialization is about durable object state.")
    if topic == "streams":
        return Note(topic, name,
                    "This class teaches stream pipelines for transformations and aggregation.",
                    "Ask whether the data should be mapped, filtered, reduced, or grouped.",
                    "A naive route would write explicit loops for every transformation step.",
                    "That is fine functionally but hides the pipeline shape.",
                    "Use streams to express the pipeline declaratively and keep the transformation stages visible.",
                    "Complexity is usually the same as the equivalent loop; the difference is readability and composition.",
                    "The lesson is that streams are a data pipeline abstraction, not magic performance.",
                    "map/filter/reduce pipeline and functional style.",
                    "source -> stream stages -> terminal operation",
                    f"{normalize(name)}: streams are about expressing transformations cleanly.")
    if topic == "consumer_supplier":
        return Note(topic, name,
                    "This class teaches functional interfaces and callback-style data flow.",
                    "Identify whether the code is producing data, consuming data, or both.",
                    "A naive route would hard-code the dependency directly into the caller.",
                    "That makes the code less reusable and harder to test.",
                    "Use consumer/supplier abstractions to decouple the producer from the consumer.",
                    "Complexity depends on the supplied action, but the structural cost is small.",
                    "The lesson is to separate data generation from data use.",
                    "producer/consumer separation and callback composition.",
                    "supplier -> value -> consumer",
                    f"{normalize(name)}: the abstraction is the direction of data flow.")
    if topic == "AnonymousClass":
        return Note(topic, name,
                    "This class teaches anonymous implementations for short-lived behavior.",
                    "Use it when a small override or callback is needed once.",
                    "A naive route would create a full named class for a tiny one-off behaviour.",
                    "That adds ceremony for very little benefit.",
                    "Use an anonymous class when the implementation is local and temporary.",
                    "Runtime cost is similar to the equivalent class; the benefit is code locality.",
                    "The lesson is to keep small behaviour close to where it is used.",
                    "callback, local override, and short-lived implementation.",
                    "interface -> anonymous implementation -> immediate use",
                    f"{normalize(name)}: anonymous classes are local behavior containers.")
    if topic == "Queue":
        return Note(topic, name,
                    "This class teaches basic FIFO queue behaviour and array/list backing ideas.",
                    "Think in terms of enqueue at one end and dequeue at the other.",
                    "A naive route would shift elements on every removal.",
                    "That is O(n) per pop and wastes work.",
                    "Use a queue structure that maintains front and rear indices or nodes.",
                    "Enqueue and dequeue are typically O(1).",
                    "The lesson is that queues model first-in-first-out workflows.",
                    "FIFO ordering and operational endpoints.",
                    "enqueue -> queue state -> dequeue",
                    f"{normalize(name)}: queues are about ordering, not sorting.")
    if topic == "Number_Theory":
        return Note(topic, name,
                    "This class teaches digit-level arithmetic and simple number transformations.",
                    "Ask whether repeated division, modular arithmetic, or digit sum rules solve the task directly.",
                    "A naive route would convert back and forth between string and integer repeatedly.",
                    "That is unnecessary when the arithmetic can be done numerically.",
                    "Use modulo and division, or a known number-theory identity like digital root.",
                    "Usually O(log n) in the number of digits.",
                    "The lesson is to recognize digit invariants and avoid full expansion.",
                    "digits, modulo, and arithmetic invariants.",
                    "number -> repeated digit math -> answer",
                    f"{normalize(name)}: number theory problems often collapse to digit invariants.")
    if topic == "parking_lot":
        return Note(topic, name,
                    "This class teaches low-level object modelling for a parking-lot style system.",
                    "Identify the entities, their relationships, and the state transitions they need.",
                    "A naive route would hard-code everything into one large procedural block.",
                    "That makes the design hard to extend when new vehicle types or rules appear.",
                    "Model the problem with classes for lot, slot, vehicle, and ticket state.",
                    "Algorithmic complexity is less important than clean composition and invariants.",
                    "The lesson is to think in terms of domain objects and responsibilities.",
                    "entities, slots, tickets, and allocation rules.",
                    "vehicle -> allocate slot -> issue ticket -> exit/update",
                    f"{normalize(name)}: system design starts with entity boundaries.")
    if topic == "main":
        return Note(topic, name,
                    "This class is part of mixed Java practice and usually demonstrates language features or small interview exercises.",
                    "Read the file as an experiment: identify what Java concept it is trying to illustrate.",
                    "A naive route would only memorize syntax without understanding the behaviour.",
                    "That makes it hard to transfer the idea to a new problem.",
                    "Focus on the concept being demonstrated, such as inheritance, recursion, or basic data flow.",
                    "Complexity depends on the specific demo, but the lesson is conceptual rather than algorithmic.",
                    "The lesson is to extract the pattern behind the example, not the example itself.",
                    "language feature, demo pattern, and conceptual behaviour.",
                    "example -> observe behaviour -> infer concept",
                    f"{normalize(name)}: this folder is about Java practice patterns and demos.")
    return generic_note(topic, name)


def classify(folder: str, file_path: Path) -> Note:
    name = file_path.stem
    topic = folder
    if topic == "heap":
        return classify_heap(name)
    if topic == "binaryTree":
        return classify_binary_tree(name)
    if topic == "graph":
        return classify_graph(name)
    if topic == "stack":
        return classify_stack(name)
    if topic == "linkedList":
        return classify_linked_list(name)
    if topic in {"arrays", "leetCode"}:
        return classify_arrays_like(topic, name)
    if topic == "dynamic":
        return classify_dynamic(name)
    if topic == "trie":
        return classify_trie(name)
    if topic == "priorityQueue":
        return classify_priority_queue(name)
    if topic == "Thread":
        return classify_thread(name)
    return classify_basic(topic, name)


def scan_topics() -> Dict[str, List[Path]]:
    topics: Dict[str, List[Path]] = defaultdict(list)
    for item in ROOT.iterdir():
        if not item.is_dir():
            continue
        if item.name in {".idea", "heap-notes"}:
            continue
        java_files = sorted(item.glob("*.java"))
        if java_files:
            topics[item.name] = java_files
    return dict(sorted(topics.items()))


def set_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.size = Pt(10)


def ensure_styles(doc: Document) -> None:
    if "Diagram" not in doc.styles:
        style = doc.styles.add_style("Diagram", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9)
    if "SectionLead" not in doc.styles:
        style = doc.styles.add_style("SectionLead", WD_STYLE_TYPE.PARAGRAPH)
        style.font.bold = True
        style.font.size = Pt(10)


def add_lead(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="SectionLead")
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(text)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_diagram(doc: Document, diagram: str) -> None:
    p = doc.add_paragraph(style="Diagram")
    p.add_run(diagram)


def add_problem_section(doc: Document, note: Note) -> None:
    doc.add_heading(note.class_name, level=2)
    add_body(doc, note.overview)
    add_lead(doc, "What this problem teaches", note.teaching_point)
    add_lead(doc, "First instinct", note.intuition)
    add_lead(doc, "Brute force", note.brute_force)
    add_lead(doc, "Why that stalls", note.bottleneck)
    add_lead(doc, "Refinement", note.better_solution)
    add_lead(doc, "Complexity", note.complexity)
    add_lead(doc, "How to recognize this pattern", note.indicators)
    add_lead(doc, "Quick takeaway", note.refresher)
    add_diagram(doc, f"Diagram: {note.diagram}")


def add_topic_header(doc: Document, topic: str, count: int) -> None:
    doc.add_heading(f"{topic} ({count} classes)", level=1)
    doc.add_paragraph(topic_overview(topic))
    doc.add_paragraph(topic_diagram(topic)).style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else doc.styles["Normal"]


def add_note_table(doc: Document, notes: List[Note], quick: bool = False) -> None:
    headers = ["Class", "Overview", "Intuition", "Brute force", "Bottleneck", "Better path", "Complexity", "Teaching point", "Indicators", "Refresher"]
    cols = headers if not quick else ["Class", "Problem -> Identify -> Solution", "Complexity", "Teaching point"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, h in enumerate(cols):
        table.rows[0].cells[i].text = h
    for note in notes:
        cells = table.add_row().cells
        if quick:
            cells[0].text = note.class_name
            cells[1].text = note.refresher
            cells[2].text = note.complexity
            cells[3].text = note.teaching_point
        else:
            vals = [
                note.class_name,
                note.overview,
                note.intuition,
                note.brute_force,
                note.bottleneck,
                note.better_solution,
                note.complexity,
                note.teaching_point,
                note.indicators,
                note.refresher,
            ]
            for idx, val in enumerate(vals):
                cells[idx].text = val


def add_detailed_topic(doc: Document, topic: str, files: List[Path]) -> None:
    doc.add_heading(f"{topic}", level=1)
    doc.add_paragraph(topic_overview(topic))
    add_diagram(doc, f"Topic flow: {topic_diagram(topic)}")
    notes = [classify(topic, fp) for fp in files]
    for note in notes:
        add_problem_section(doc, note)
        doc.add_paragraph("")
    doc.add_paragraph("")


def add_summary_topic(doc: Document, topic: str, files: List[Path]) -> None:
    doc.add_heading(f"{topic}", level=1)
    doc.add_paragraph(topic_overview(topic))
    notes = [classify(topic, fp) for fp in files]
    add_note_table(doc, notes, quick=False)
    doc.add_paragraph("")


def add_quick_topic(doc: Document, topic: str, files: List[Path]) -> None:
    doc.add_heading(f"{topic}", level=1)
    doc.add_paragraph(topic_overview(topic))
    notes = [classify(topic, fp) for fp in files]
    add_note_table(doc, notes, quick=True)
    doc.add_paragraph("")


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def write_inventory_doc(topics: Dict[str, List[Path]]) -> Path:
    doc = Document()
    set_document_style(doc)
    add_title(doc, "Java Topic Inventory", "Folders under src that contain topic-style Java classes")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "Class count"
    table.rows[0].cells[2].text = "Theme"
    for topic, files in topics.items():
        row = table.add_row().cells
        row[0].text = topic
        row[1].text = str(len(files))
        row[2].text = topic_overview(topic)
    path = OUT / "01-topic-inventory.docx"
    doc.save(path)
    return path


def write_detailed_doc(topics: Dict[str, List[Path]]) -> Path:
    doc = Document()
    set_document_style(doc)
    ensure_styles(doc)
    add_title(doc, "Java Study Notes", "Progressive notes with brute force, refinement, and complexity")
    doc.add_paragraph(
        "Each chapter below reads like a study guide. The structure is deliberate: understand the concept first, try the naive route, notice the bottleneck, and then build the better approach step by step."
    )
    doc.add_heading("How to read these notes", level=1)
    doc.add_paragraph("1. Start by asking what the class is really teaching, not just what code it contains.")
    doc.add_paragraph("2. Follow the brute-force path first so the optimized idea has context.")
    doc.add_paragraph("3. Pay attention to the failure point; that is usually the root lesson.")
    doc.add_paragraph("4. Use the diagram and complexity notes to connect intuition with implementation.")
    for topic, files in topics.items():
        add_detailed_topic(doc, topic, files)
        doc.add_page_break()
    path = OUT / "02-detailed-notes.docx"
    doc.save(path)
    return path


def write_summary_doc(topics: Dict[str, List[Path]]) -> Path:
    doc = Document()
    set_document_style(doc)
    add_title(doc, "Java Combined Notes", "Table view of problem, intuition, concepts, and intent")
    for section in doc.sections:
        set_landscape(section)
    doc.add_paragraph("This document collapses each class into a compact table row so the study path stays easy to scan.")
    for topic, files in topics.items():
        doc.add_heading(topic, level=1)
        doc.add_paragraph(topic_overview(topic))
        notes = [classify(topic, fp) for fp in files]
        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        headers = ["Problem", "Overview", "Intuition", "Concepts used", "How to identify", "How to tackle", "Indicators", "Key takeaway", "Underlying intent"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for note in notes:
            row = table.add_row().cells
            row[0].text = note.class_name
            row[1].text = note.overview
            row[2].text = note.intuition
            row[3].text = topic_overview(note.topic)
            row[4].text = note.indicators
            row[5].text = note.better_solution
            row[6].text = note.indicators
            row[7].text = note.teaching_point
            row[8].text = note.refresher
        doc.add_page_break()
    path = OUT / "03-combined-summary.docx"
    doc.save(path)
    return path


def write_quick_doc(topics: Dict[str, List[Path]]) -> Path:
    doc = Document()
    set_document_style(doc)
    add_title(doc, "Java Quick Refresher", "3-4 line problem -> identify -> solution prompts")
    doc.add_paragraph("Use this as a fast review sheet before interviews or revision sessions.")
    for topic, files in topics.items():
        doc.add_heading(topic, level=1)
        doc.add_paragraph(topic_overview(topic))
        notes = [classify(topic, fp) for fp in files]
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Problem"
        table.rows[0].cells[1].text = "Identify -> solve"
        table.rows[0].cells[2].text = "Mental model"
        for note in notes:
            row = table.add_row().cells
            row[0].text = note.class_name
            row[1].text = note.refresher
            row[2].text = note.teaching_point
        doc.add_page_break()
    path = OUT / "04-quick-refresher.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    topics = scan_topics()
    if not topics:
        raise SystemExit("No topic folders found under src")
    inventory = write_inventory_doc(topics)
    detailed = write_detailed_doc(topics)
    summary = write_summary_doc(topics)
    quick = write_quick_doc(topics)
    print("\n".join(str(p) for p in [inventory, detailed, summary, quick]))


if __name__ == "__main__":
    main()
